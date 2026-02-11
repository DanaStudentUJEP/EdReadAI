import os
import io
import json
import re
import requests
import streamlit as st
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE


# =========================
# OpenAI
# =========================
OPENAI_CHAT_URL = "https://api.openai.com/v1/chat/completions"


def get_openai_key() -> str:
    if hasattr(st, "secrets") and "OPENAI_API_KEY" in st.secrets:
        return str(st.secrets["OPENAI_API_KEY"]).strip()
    return (os.getenv("OPENAI_API_KEY") or "").strip()


def get_openai_model() -> str:
    if hasattr(st, "secrets") and "OPENAI_MODEL" in st.secrets:
        return str(st.secrets["OPENAI_MODEL"]).strip()
    return (os.getenv("OPENAI_MODEL") or "gpt-4o-mini").strip()


def call_openai_chat(system_prompt: str, user_prompt: str, temperature: float = 0.2, max_tokens: int = 2200) -> str:
    api_key = get_openai_key()
    if not api_key:
        raise RuntimeError("Chybí OPENAI_API_KEY (Streamlit Cloud → Settings → Secrets).")

    payload = {
        "model": get_openai_model(),
        "temperature": float(temperature),
        "max_tokens": int(max_tokens),
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    r = requests.post(OPENAI_CHAT_URL, headers=headers, json=payload, timeout=90)

    if r.status_code != 200:
        raise RuntimeError(f"OpenAI API chyba ({r.status_code}): {r.text}")

    data = r.json()
    return data["choices"][0]["message"]["content"]


# =========================
# DOCX helpers
# =========================
def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def add_spacer(doc: Document, cm: float = 0.2) -> None:
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(int(cm * 28.35))


def doc_to_bytes(doc: Document) -> bytes:
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def safe_filename(name: str) -> str:
    # bezpečné jméno souboru pro Windows
    name = re.sub(r"[\\/:*?\"<>|]+", "", name)
    name = name.strip()
    return name if name else "edread_ai"


def asset_candidates() -> Dict[str, List[str]]:
    """
    Více názvů pro stejné tabulky – aby to sedělo na různé verze souborů.
    Ulož do assets/ aspoň jednu z uvedených variant.
    """
    return {
        "karetni_table": [
            "assets/karetni_table.png",
            "assets/karetni_table_only.png",
        ],
        "sladke_table": [
            "assets/sladke_table.png",
            "assets/sladke_p1.png",
            "assets/sladke_p1_300.png",
        ],
        "venecky_table": [
            "assets/venecky_table.png",
            "assets/venecky_p2_300.png",
        ],
    }


def find_existing_asset(paths: List[str]) -> Optional[str]:
    for p in paths:
        if os.path.exists(p):
            return p
    return None


def add_image_if_exists(doc: Document, path: str, width_cm: float = 16.0, center: bool = True) -> bool:
    if not path or not os.path.exists(path):
        return False
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return True


# =========================
# Detekce „speciálních“ textů
# =========================
def detect_pack(title: str, full_text: str) -> str:
    t = (title or "").lower()
    x = (full_text or "").lower()
    if "karetní hra" in t or "karetni hra" in t or "karetní" in t or "karetni" in t:
        return "karetni"
    if "sladké mámení" in t or "sladke mamen" in t or "mámení" in t or "mamen" in t:
        return "sladke"
    if "věnečky" in t or "venecky" in t:
        return "venecky"

    # fallback podle obsahu
    if "kdo přebije koho" in x or "žolík" in x or "chameleon" in x:
        return "karetni"
    if "věneček" in x and "cukrárn" in x:
        return "venecky"
    if "mámení" in x and "sladké" in x:
        return "sladke"

    return "custom"


# =========================
# Karetní hra – zvířata a pyramida
# =========================
ANIMALS: List[Tuple[str, str]] = [
    ("🦟", "komár"),
    ("🐭", "myš"),
    ("🐟", "sardinka"),
    ("🦔", "ježek"),
    ("🐟", "okoun"),
    ("🦊", "liška"),
    ("🦭", "tuleň"),
    ("🦁", "lev"),
    ("🐻‍❄️", "lední medvěd"),
    ("🐊", "krokodýl"),
    ("🐘", "slon"),
    ("🐬", "kosatka"),
    ("🦎", "chameleon (žolík)"),
]

# Logika pyramidy „nejvyšší = nejsilnější“ – upravuješ jen pořadí.
# (Když máš v pravidlech přesné pořadí, sem ho dej 1:1.)
PYRAMID_ORDER_STRONG_TO_WEAK = [
    "kosatka",
    "slon",
    "krokodýl",
    "lední medvěd",
    "lev",
    "tuleň",
    "liška",
    "okoun",
    "ježek",
    "sardinka",
    "myš",
    "komár",
    "chameleon (žolík)",  # žolík můžeš mít kde chceš – pokud má být jinak, přesuň ho
]


def add_karetni_cards_3col(doc: Document) -> None:
    add_h2(doc, "Kartičky zvířat (vystřihni)")
    doc.add_paragraph("Vystřihni kartičky. Pak je použiješ do pyramidy síly.")
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # rozdělení do 3 sloupců
    cells = []
    for emoji, name in ANIMALS:
        cells.append((emoji, name))

    # doplnění do řádků
    idx = 0
    while idx < len(cells):
        row_cells = table.add_row().cells
        for c in range(3):
            if idx < len(cells):
                emoji, name = cells[idx]
                p1 = row_cells[c].paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run(emoji)
                r1.font.size = Pt(26)

                p2 = row_cells[c].add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(name)
                r2.bold = True
                r2.font.size = Pt(12)
                idx += 1
            else:
                row_cells[c].text = ""

    # trochu prostoru
    doc.add_paragraph("")


def add_pyramid_column(doc: Document, card_width_cm: float = 6.0, box_height_cm: float = 1.6) -> None:
    """
    Sloupcová „pyramida“ – jedno okénko na každé zvíře.
    Okénka jsou úmyslně větší, aby se do nich vešly kartičky.
    """
    add_h2(doc, "Pyramida síly (nalep kartičky)")
    doc.add_paragraph(
        "Nalep kartičky do pyramidy podle pravidel hry: Nahoře nejsilnější, dole nejslabší."
    )

    # 2 sloupce: vlevo pořadí (1–13), vpravo okénko pro kartičku
    table = doc.add_table(rows=len(PYRAMID_ORDER_STRONG_TO_WEAK), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, animal_name in enumerate(PYRAMID_ORDER_STRONG_TO_WEAK, start=1):
        row = table.rows[i - 1]
        row.height = Cm(box_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        left = row.cells[0]
        right = row.cells[1]

        left.text = f"{i}."
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # okénko – necháme prázdné, ale doplníme jemný popisek (učitel může vypnout)
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(" ")  # prázdno, aby se držela výška

        # šířky sloupců (python-docx neumí 100% fixně, ale Word to drží dobře)
        left.width = Cm(1.0)
        right.width = Cm(card_width_cm)

    doc.add_paragraph("")


def add_karetni_pack_extras(doc: Document, include_table: bool = True) -> None:
    """
    Přidá do pracovního listu Karetní hry:
    - tabulku „Kdo přebije koho?“ (PNG)
    - pyramidu + kartičky
    """
    add_h2(doc, "Pomůcky k pravidlům hry")
    # tabulka (PNG) – musí být v assets
    if include_table:
        pth = find_existing_asset(asset_candidates()["karetni_table"])
        if pth:
            doc.add_paragraph("Tabulka: Kdo přebije koho?")
            add_image_if_exists(doc, pth, width_cm=16.0, center=True)
        else:
            doc.add_paragraph("⚠️ Tabulka „Kdo přebije koho?“ nebyla nalezena (chybí PNG v assets/).")

    add_spacer(doc, 0.15)
    add_pyramid_column(doc, card_width_cm=6.5, box_height_cm=1.7)
    add_karetni_cards_3col(doc)


# =========================
# AI – struktura z vlastního textu
# =========================
@dataclass
class GeneratedStructure:
    simpl: str
    lmp: str
    drama_intro: str
    drama_scene: List[Tuple[str, str]]
    glossary: Dict[str, str]
    questions_A: List[str]
    questions_B: List[str]
    questions_C: List[str]


def ai_generate_structure(full_text: str, grade: int, title: str) -> GeneratedStructure:
    """
    Z jednoho vstupního textu vygeneruje:
    - zjednodušenou verzi
    - LMP/SPU verzi
    - dramatizaci (intro + 3–6 replik)
    - slovníček pojmů
    - otázky A/B/C
    """
    if not get_openai_key():
        return GeneratedStructure(
            simpl=full_text,
            lmp=full_text,
            drama_intro="(Dramatizace není k dispozici – chybí OPENAI_API_KEY.)",
            drama_scene=[],
            glossary={},
            questions_A=["(Otázky A nejsou k dispozici – chybí OPENAI_API_KEY.)"],
            questions_B=["(Otázky B nejsou k dispozici – chybí OPENAI_API_KEY.)"],
            questions_C=["(Otázky C nejsou k dispozici – chybí OPENAI_API_KEY.)"],
        )

    system = (
        "Jsi odborník na český jazyk, čtenářskou gramotnost a RVP ZV. "
        "Umíš tvořit pracovní listy ve stylu ČŠI (čtení s porozuměním). "
        "Výstup musí být validní JSON, žádný komentář navíc."
    )

    user = f"""
Máš vytvořit pracovní list pro žáky {grade}. ročníku ZŠ.
Název úlohy: {title}

Vstupní text (plná verze):
\"\"\"{full_text}\"\"\"

ÚKOL:
1) Vytvoř ZJEDNODUŠENOU verzi textu (pro běžné žáky).
2) Vytvoř LMP/SPU verzi (velmi krátké věty, maximální srozumitelnost).
3) Vytvoř krátkou DRAMATIZACI:
   - 1–2 věty „drama_intro“ (co se bude hrát, proč).
   - 3–6 replik ve formátu: [ ["Role", "replika"], ... ]
   - Scénka má být „bez pomůcek“, jen hraní rolí.
4) Vytvoř SLOVNÍČEK:
   - vyber 8–14 slov z textu, která mohou být pro žáky obtížná,
   - ke každému napiš krátké vysvětlení (max 12 slov),
   - vrať jako slovník {{ "slovo": "vysvětlení" }}.
5) Vytvoř OTÁZKY A/B/C:
   - A: 3–4 otázky na vyhledávání informací.
   - B: 2–3 otázky na porozumění a interpretaci.
   - C: 2–3 otázky na názor / kritické čtení (žák zdůvodní).

VRAŤ POUZE JSON VE FORMÁTU:

{{
  "simpl": "...",
  "lmp": "...",
  "drama_intro": "...",
  "drama_scene": [
    ["Role 1", "replika 1"],
    ["Role 2", "replika 2"]
  ],
  "glossary": {{
    "slovo1": "vysvětlení1",
    "slovo2": "vysvětlení2"
  }},
  "questions_A": ["otázka A1", "otázka A2"],
  "questions_B": ["otázka B1", "otázka B2"],
  "questions_C": ["otázka C1", "otázka C2"]
}}
"""

    out = call_openai_chat(system, user, temperature=0.2, max_tokens=2600)
    data = json.loads(out)

    simpl = str(data.get("simpl", full_text)).strip() or full_text
    lmp = str(data.get("lmp", full_text)).strip() or full_text
    drama_intro = str(data.get("drama_intro", "")).strip()

    drama_scene_raw = data.get("drama_scene", [])
    drama_scene: List[Tuple[str, str]] = []
    if isinstance(drama_scene_raw, list):
        for item in drama_scene_raw:
            if isinstance(item, (list, tuple)) and len(item) == 2:
                role = str(item[0]).strip()
                line = str(item[1]).strip()
                if role and line:
                    drama_scene.append((role, line))

    glossary_raw = data.get("glossary", {})
    glossary: Dict[str, str] = {}
    if isinstance(glossary_raw, dict):
        for k, v in glossary_raw.items():
            kk = str(k).strip()
            vv = str(v).strip()
            if kk and vv:
                glossary[kk] = vv

    def _clean_list(key: str) -> List[str]:
        arr = data.get(key, [])
        out_list: List[str] = []
        if isinstance(arr, list):
            for q in arr:
                qq = str(q).strip()
                if qq:
                    out_list.append(qq)
        return out_list or [f"(Žádné otázky v sekci {key} – zkus generovat znovu.)"]

    questions_A = _clean_list("questions_A")
    questions_B = _clean_list("questions_B")
    questions_C = _clean_list("questions_C")

    return GeneratedStructure(
        simpl=simpl,
        lmp=lmp,
        drama_intro=drama_intro or "Na začátku si krátce zahrajeme scénku, která ti pomůže pochopit, o čem text bude.",
        drama_scene=drama_scene,
        glossary=glossary,
        questions_A=questions_A,
        questions_B=questions_B,
        questions_C=questions_C,
    )


# =========================
# DOCX – pracovní list
# =========================
def add_glossary_block(doc: Document, glossary: Dict[str, str]) -> None:
    add_h2(doc, "Slovníček pojmů (na závěr)")
    if not glossary:
        doc.add_paragraph("Slovníček není k dispozici.")
        return

    doc.add_paragraph("Nejdřív si slovíčka projděte společně s učitelem/kou. Pak se vraťte k textu.")
    for w, expl in glossary.items():
        p = doc.add_paragraph()
        r = p.add_run(f"• {w} — ")
        r.bold = True
        p.add_run(expl)
        p.add_run("  | Poznámka: ________________________________")


def add_tables_for_pack_inside_text(doc: Document, pack: str) -> None:
    """
    Vloží tabulku/tabulky jako obrázek do části „Text pro čtení“.
    Tabulky jsou nutné i pro zjednodušenou a LMP verzi.
    """
    ac = asset_candidates()

    if pack == "karetni":
        pth = find_existing_asset(ac["karetni_table"])
        if pth:
            doc.add_paragraph("Tabulka z pravidel: Kdo přebije koho?")
            add_image_if_exists(doc, pth, width_cm=16.0, center=True)
        else:
            doc.add_paragraph("⚠️ Chybí tabulka (PNG) pro Karetní hru v assets/.")

    elif pack == "sladke":
        pth = find_existing_asset(ac["sladke_table"])
        if pth:
            doc.add_paragraph("Tabulka z textu (pro práci s otázkami):")
            add_image_if_exists(doc, pth, width_cm=16.0, center=True)
        else:
            doc.add_paragraph("⚠️ Chybí tabulka (PNG) pro Sladké mámení v assets/.")

    elif pack == "venecky":
        pth = find_existing_asset(ac["venecky_table"])
        if pth:
            doc.add_paragraph("Tabulka z textu (pro práci s otázkami):")
            add_image_if_exists(doc, pth, width_cm=16.0, center=True)
        else:
            doc.add_paragraph("⚠️ Chybí tabulka (PNG) pro Věnečky v assets/.")


def build_student_doc(
    title: str,
    grade: int,
    variant_label: str,
    text_variant: str,
    drama_intro: str,
    drama_scene: List[Tuple[str, str]],
    glossary: Dict[str, str],
    questions_A: List[str],
    questions_B: List[str],
    questions_C: List[str],
    pack: str,
) -> Document:
    doc = Document()
    set_doc_defaults(doc)

    add_h1(doc, f"NÁZEV ÚLOHY: {title} — {variant_label}")
    doc.add_paragraph(f"Ročník: {grade}. třída")
    doc.add_paragraph("JMÉNO: ________________________________    DATUM: _______________")
    add_spacer(doc, 0.2)

    # 1) dramatizace
    add_h2(doc, "1) Úvod a krátká dramatizace (začátek hodiny)")
    doc.add_paragraph(
        "Nejdřív si zahrajeme krátkou scénku. Pomůže ti rychle pochopit, o čem text bude."
    )
    doc.add_paragraph(drama_intro)
    for role, line in drama_scene:
        doc.add_paragraph(f"{role}: {line}")
    add_spacer(doc, 0.2)

    # 2) text + tabulky uvnitř textu
    add_h2(doc, "2) Text pro čtení")
    doc.add_paragraph(text_variant)
    add_spacer(doc, 0.15)
    # tabulky nutné pro odpovědi – ve všech verzích
    if pack in ("karetni", "sladke", "venecky"):
        add_tables_for_pack_inside_text(doc, pack)
        add_spacer(doc, 0.2)

    # 2b) Karetní hra – pomůcky (pyramida + kartičky + tabulka)
    if pack == "karetni":
        add_karetni_pack_extras(doc, include_table=False)  # tabulka už je vložená u textu
        add_spacer(doc, 0.2)

    # 3) otázky
    add_h2(doc, "3) Otázky k textu")

    doc.add_paragraph("A) Najdi v textu (vyhledávání informací):")
    for q in questions_A:
        doc.add_paragraph(f"• {q}\n  Odpověď: ______________________________________________")

    add_spacer(doc, 0.15)
    doc.add_paragraph("B) Přemýšlej a vysvětli (porozumění / interpretace):")
    for q in questions_B:
        doc.add_paragraph(
            f"• {q}\n  Odpověď: ______________________________________________\n  ______________________________________________"
        )

    add_spacer(doc, 0.15)
    doc.add_paragraph("C) Můj názor (kritické čtení / argumentace):")
    for q in questions_C:
        doc.add_paragraph(
            f"• {q}\n  Odpověď: ______________________________________________\n  ______________________________________________"
        )

    add_spacer(doc, 0.25)
    # slovníček až na konci
    add_glossary_block(doc, glossary)

    return doc


def build_method_doc(
    title: str,
    grade: int,
    full_text: str,
    structure: GeneratedStructure,
    pack: str,
) -> Document:
    doc = Document()
    set_doc_defaults(doc)

    add_h1(doc, f"Metodický list pro učitele — {title}")
    doc.add_paragraph(f"Ročník: {grade}. třída")

    add_h2(doc, "Cíl hodiny")
    doc.add_paragraph(
        "Rozvoj čtenářské gramotnosti v souladu s RVP ZV: vyhledávání informací, porozumění textu, interpretace, "
        "kritické čtení a formulace vlastního názoru."
    )

    add_h2(doc, "Doporučený postup (45 min)")
    doc.add_paragraph("1) Úvod + dramatizace (5–7 min) – krátká scénka z pracovního listu, motivace.")
    doc.add_paragraph(
        "2) Slovníček (5–8 min) – i když je na konci listu, pracujte s ním hned po dramatizaci. "
        "Vyberte slova, která mohou brzdit porozumění; žáci si doplní poznámky."
    )
    doc.add_paragraph("3) Čtení textu (10–12 min) – tiché čtení / čtení po odstavcích.")
    doc.add_paragraph(
        "4) Otázky A/B/C (15–18 min) – A: dohledání informace, B: vysvětlení vlastními slovy, "
        "C: názor + zdůvodnění."
    )
    doc.add_paragraph("5) Reflexe (2–3 min) – co pomohlo porozumět (dramatizace, tabulka, slovníček).")

    add_h2(doc, "Tabulky / opory v textu")
    if pack in ("karetni", "sladke", "venecky"):
        doc.add_paragraph("Tabulka z původního textu je vložená přímo v části „Text pro čtení“ ve všech verzích.")
    if pack == "karetni":
        doc.add_paragraph("Karetní hra: navíc je přiložená pyramida síly a kartičky zvířat (vystřižení a lepení).")

    add_h2(doc, "Poznámka k verzím")
    doc.add_paragraph("Plná verze: původní text (vstup učitele).")
    doc.add_paragraph("Zjednodušená verze: kratší věty, jednodušší slovní zásoba, zachování klíčových informací.")
    doc.add_paragraph("LMP/SPU verze: velmi krátké věty, maximální srozumitelnost, odstranění složitých souvětí.")
    doc.add_paragraph(
        "Rozdíly mezi verzemi jsou pouze v textu (plný / zjednodušený / LMP). "
        "Tabulky zůstávají ve všech verzích stejné, aby šly vypracovat otázky."
    )

    add_h2(doc, "Vstupní text (plná verze)")
    doc.add_paragraph(full_text)

    add_h2(doc, "Zjednodušená verze (náhled)")
    doc.add_paragraph(structure.simpl)

    add_h2(doc, "LMP/SPU verze (náhled)")
    doc.add_paragraph(structure.lmp)

    return doc


# =========================
# Generování všech dokumentů
# =========================
def generate_all_from_text(title: str, grade: int, full_text: str) -> Dict[str, bytes]:
    pack = detect_pack(title, full_text)
    structure = ai_generate_structure(full_text, grade, title)

    doc_full = build_student_doc(
        title=title,
        grade=grade,
        variant_label="PLNÝ",
        text_variant=full_text,
        drama_intro=structure.drama_intro,
        drama_scene=structure.drama_scene,
        glossary=structure.glossary,
        questions_A=structure.questions_A,
        questions_B=structure.questions_B,
        questions_C=structure.questions_C,
        pack=pack,
    )

    doc_simpl = build_student_doc(
        title=title,
        grade=grade,
        variant_label="ZJEDNODUŠENÝ",
        text_variant=structure.simpl,
        drama_intro=structure.drama_intro,
        drama_scene=structure.drama_scene,
        glossary=structure.glossary,
        questions_A=structure.questions_A,
        questions_B=structure.questions_B,
        questions_C=structure.questions_C,
        pack=pack,
    )

    doc_lmp = build_student_doc(
        title=title,
        grade=grade,
        variant_label="LMP/SPU",
        text_variant=structure.lmp,
        drama_intro=structure.drama_intro,
        drama_scene=structure.drama_scene,
        glossary=structure.glossary,
        questions_A=structure.questions_A,
        questions_B=structure.questions_B,
        questions_C=structure.questions_C,
        pack=pack,
    )

    doc_method = build_method_doc(
        title=title,
        grade=grade,
        full_text=full_text,
        structure=structure,
        pack=pack,
    )

    return {
        "pl_full": doc_to_bytes(doc_full),
        "pl_simpl": doc_to_bytes(doc_simpl),
        "pl_lmp": doc_to_bytes(doc_lmp),
        "method": doc_to_bytes(doc_method),
    }


# =========================
# Streamlit state + UI
# =========================
def ensure_state():
    if "files" not in st.session_state:
        st.session_state["files"] = {}
    if "names" not in st.session_state:
        st.session_state["names"] = {}
    if "generated" not in st.session_state:
        st.session_state["generated"] = False


def show_downloads():
    files: Dict[str, bytes] = st.session_state.get("files", {})
    names: Dict[str, str] = st.session_state.get("names", {})
    if not files:
        return

    st.subheader("Stažení dokumentů")

    labels = {
        "pl_full": "⬇️ Pracovní list – plná verze",
        "pl_simpl": "⬇️ Pracovní list – zjednodušená verze",
        "pl_lmp": "⬇️ Pracovní list – LMP/SPU verze",
        "method": "⬇️ Metodický list pro učitele",
    }

    order = ["pl_full", "pl_simpl", "pl_lmp", "method"]
    cols = st.columns(2)
    for i, k in enumerate(order):
        if k in files:
            with cols[i % 2]:
                st.download_button(
                    label=labels.get(k, f"Stáhnout {k}"),
                    data=files[k],
                    file_name=names.get(k, f"{k}.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{k}",
                )

    if st.button("🧹 Vymazat vygenerované soubory", key="clear_btn"):
        st.session_state["files"] = {}
        st.session_state["names"] = {}
        st.session_state["generated"] = False
        st.success("Vygenerované soubory byly vymazány.")


def main():
    st.set_page_config(page_title="EdRead AI – vlastní text", layout="centered")
    ensure_state()

    st.title("EdRead AI — pracovní list z vlastního textu")

    if get_openai_key():
        st.success(f"OPENAI_API_KEY nalezen. Model: {get_openai_model()}")
    else:
        st.warning("Chybí OPENAI_API_KEY → vše poběží v nouzovém režimu (bez AI úprav).")

    st.info(
        "Vlož vlastní text. EdRead AI z něj vytvoří plný, zjednodušený a LMP/SPU pracovní list "
        "s dramatizací, slovníčkem a otázkami A/B/C. "
        "Pro texty Karetní hra / Sladké mámení / Věnečky navíc vloží tabulky (PNG z assets/) a u Karetní hry i pyramidu + kartičky."
    )

    title = st.text_input("Název úlohy:", value="Moje čtení s porozuměním")
    grade = st.number_input("Ročník (1–9):", min_value=1, max_value=9, value=5, step=1)
    full_text = st.text_area("Vlož text pro čtení:", height=320, placeholder="Sem vlož celý text, se kterým chceš pracovat...")

    # rychlá kontrola assets
    with st.expander("🔎 Kontrola tabulek v assets/ (doporučeno)", expanded=False):
        ac = asset_candidates()
        for key, candidates in ac.items():
            found = find_existing_asset(candidates)
            if found:
                st.success(f"{key}: nalezeno → {found}")
            else:
                st.warning(f"{key}: nenalezeno (nahraj PNG do assets/)")

    if st.button("Vygenerovat pracovní listy", type="primary", key="btn_generate"):
        if not full_text.strip():
            st.error("Nejdřív vlož text.")
        else:
            try:
                with st.spinner("Generuji pracovní listy…"):
                    out = generate_all_from_text(title, int(grade), full_text.strip())

                base = safe_filename(title)
                st.session_state["files"] = out
                st.session_state["names"] = {
                    "pl_full": f"pracovni_list_{base}_plny.docx",
                    "pl_simpl": f"pracovni_list_{base}_zjednoduseny.docx",
                    "pl_lmp": f"pracovni_list_{base}_LMP_SPU.docx",
                    "method": f"metodika_{base}.docx",
                }
                st.session_state["generated"] = True
                st.success("Hotovo. Dokumenty jsou připravené ke stažení.")
            except Exception as e:
                st.error(f"Došlo k chybě při generování: {e}")

    # Tlačítka zůstanou – držíme bytes v session_state
    show_downloads()


if __name__ == "__main__":
    main()
